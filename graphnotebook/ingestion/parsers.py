"""
Document parsers. Extract raw text + metadata from PDF, DOCX, MD/TXT.
All parsers return a standardized ParsedDocument dataclass.
"""

import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument


@dataclass
class PageContent:
    page_number: int
    text: str
    section_header: Optional[str] = None


@dataclass
class ParsedDocument:
    filename: str
    file_type: str
    file_hash: str
    pages: List[PageContent]
    raw_text: str  # full concatenated text
    raw_text_length: int
    metadata: dict = field(default_factory=dict)


def parse_pdf(file_path: str) -> ParsedDocument:
    """Extract text from PDF with page-level granularity."""
    path = Path(file_path)
    doc = fitz.open(file_path)
    pages = []
    full_text_parts = []

    for i, page in enumerate(doc):
        text = page.get_text("text")
        pages.append(PageContent(page_number=i + 1, text=text))
        full_text_parts.append(text)

    raw_text = "\n\n".join(full_text_parts)
    file_hash = hashlib.sha256(path.read_bytes()).hexdigest()

    return ParsedDocument(
        filename=path.name,
        file_type="pdf",
        file_hash=file_hash,
        pages=pages,
        raw_text=raw_text,
        raw_text_length=len(raw_text),
        metadata={"page_count": len(doc)},
    )


def parse_docx(file_path: str) -> ParsedDocument:
    """Extract text from DOCX with paragraph-level structure."""
    path = Path(file_path)
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    raw_text = "\n\n".join(paragraphs)
    file_hash = hashlib.sha256(path.read_bytes()).hexdigest()

    return ParsedDocument(
        filename=path.name,
        file_type="docx",
        file_hash=file_hash,
        pages=[PageContent(page_number=1, text=raw_text)],
        raw_text=raw_text,
        raw_text_length=len(raw_text),
    )


def parse_text(file_path: str) -> ParsedDocument:
    """Parse plain text / markdown files."""
    path = Path(file_path)
    raw_text = path.read_text(encoding="utf-8")
    file_hash = hashlib.sha256(raw_text.encode("utf-8")).hexdigest()

    return ParsedDocument(
        filename=path.name,
        file_type=path.suffix.lstrip("."),
        file_hash=file_hash,
        pages=[PageContent(page_number=1, text=raw_text)],
        raw_text=raw_text,
        raw_text_length=len(raw_text),
    )


PARSER_MAP = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".md": parse_text,
    ".txt": parse_text,
}


def parse_document(file_path: str) -> ParsedDocument:
    """Route to the correct parser based on file extension."""
    ext = Path(file_path).suffix.lower()
    parser = PARSER_MAP.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file type: {ext}")
    return parser(file_path)
